#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# Author: Shieber
# Date: 2019.07.24
# Modified: 2020.07.24

import re,sys,time
from os import listdir,getcwd
from docx import Document 
from os.path import basename, exists

def write2docx(docxname,title,txt,logname):
    '''将text内容写入docx'''
    if not txt:
        with open(logname,'w+') as logobj:
            logobj.write('Error:no texture')
        sys.exit(-1)

    if exists(docxname):
        docx = Document(docxname)
    else:
        docx = Document()

    try:
        docx.add_heading(title)
        docx.add_paragraph(txt)
    except Exception as err:
        with open(logname,'w+') as logobj:
            logobj.write(f'Error: {err}')
        sys.exit(-1)

    docx.save(docxname)

def getTitle(fl,logname):
    '''获取txt和title'''
    try:
        with open(fl) as txtobj:
            title = txtobj.readline()
            text  = txtobj.read()
        return title, text
    except Exception as err:
        with open(logname,'w+') as logobj:
            logobj.write(f'Error: {err}')
        sys.exit(-1)

def transfer(files):
    logname = ''.join([getcwd(),'/','err.log'])
    for fl in files:
        if not fl.endswith('.txt'):
            continue
        docxname  = ''.join([fl.split('.')[0],'.docx'])
        title,txt = getTitle(fl,logname)
        write2docx(docxname,title,txt,logname)

def text2docx():
    '''主函数'''
    argv = sys.argv
    if len(argv) < 2:
        program = basename(argv[0])
        print(f"Usage: {program} test.txt or {program} -a")
        sys.exit(-1)

    if '-a' == argv[1] or '--all' == argv[1]:
        fils = listdir('.')
    else:
        fils = argv[1:]

    transfer(fils)

if __name__ == "__main__":
    start = time.time()
    text2docx()
    end = time.time()
    print(f'耗时：{end - start:.2f}(s)')
