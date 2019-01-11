#!/usr/bin/python
# -*- coding: utf-8 -*-
# Copyright 2019 Shieber
# All Rights Reserved.
#
#    Licensed under the Apache License, Version 2.0 (the "License"); you may
#    not use this file except in compliance with the License. You may obtain
#    a copy of the License at
#
#         http://www.apache.org/licenses/LICENSE-2.0
#
#    Unless required by applicable law or agreed to in writing, software
#    distributed under the License is distributed on an "AS IS" BASIS, WITHOUT
#    WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied. See the
#    License for the specific language governing permissions and limitations
#    under the License.
# 将txt文件转换为docx文档
# -a 将当前目录所有txt文件转换为docx格式

from os import listdir
from os.path import  basename 
from os.path import  exists
from docx import Document 
import re 

import sys
reload(sys)
sys.setdefaultencoding('utf-8')  #解决中文解析出错的问题,decode解决写入问题

def istxt(filename):
	'''判断是否为txt文档'''
	if filename.endswith('.txt'):
		return True
	else:
		return False

def getprefix(filename):
	'''提取电子书或任意文件的前缀名'''
	name_pattern = re.compile(r'(.*?)\.')
	try:
		matched = name_pattern.search(filename)   #文件名，书名，不含点和后缀
		prefix  = matched.group(1)
	except Exception:                        
		print("Name error")
		sys.exit(-1)

	return prefix                                   

def get_titl_text(filename):
	'''获取txt和title'''
	try:
		txtobj = open(filename)
		title  = txtobj.readline().decode('utf-8')
		text   = txtobj.read().decode('utf-8')     #decode还原为utf-8  encode
		txtobj.close() 
	except Exception as err:
		print("Error, file does not exist")
		txtobj.close() 
		sys.exit(-1)

	return title, text

def write2docx(docxname,title,text):
	'''将text内容写入docx'''
	if not text:
		print("Erro,no texture")
		sys.exit(-1)

	if exists(docxname):
		docx = Document(docxname)
	else:
		docx = Document()

	try:
		docx.add_heading(title)
		docx.add_paragraph(text)
	except Exception as err:
		print("错误:%s"%err)
		sys.exit(-1)

	docx.save(docxname)

def text2docx():
	'''主函数'''
	argv = sys.argv
	if len(argv) < 2:
		prog_name = basename(argv[0])
		print("Usage: %s [s1.txt] or %s -a"%(prog_name,prog_name))
		sys.exit(-1)

	if '-a' == argv[1]:
		for fil in listdir('.'):
			if not istxt(fil):
				continue
			docxname = getprefix(fil) + '.docx'
			title, text = get_titl_text(fil)
			write2docx(docxname,title,text)
	else:
		for fil in argv[1:]:
			if not istxt(fil):
				continue
			docxname = getprefix(fil) + '.docx'
			title, text = get_titl_text(fil)
			write2docx(docxname,title,text)

if __name__ == "__main__":
	text2docx()
